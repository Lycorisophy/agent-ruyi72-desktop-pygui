#!/usr/bin/env python3
"""
Word Document Writer for Ruyi72
Creates professional .docx documents
"""

import argparse
import json
import sys
from pathlib import Path
from datetime import datetime


def create_word_document(content_dict: dict, output_path: str) -> dict:
    """
    创建 Word 文档
    
    Args:
        content_dict: 文档内容配置
        output_path: 输出文件路径
    
    Returns:
        创建结果
    """
    results = {
        'success': False,
        'path': output_path,
        'message': ''
    }
    
    try:
        # 尝试导入 python-docx
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        results['message'] = """
❌ 缺少依赖库 python-docx

请安装依赖：
pip install python-docx

或者使用 MD 格式替代：
md-writer 技能可以直接生成 Markdown 文档，支持转换为多种格式。
"""
        return results
    
    # 创建文档
    doc = Document()
    
    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)
    
    # 标题
    if 'title' in content_dict:
        title = doc.add_heading()
        title_run = title.add_run(content_dict['title'])
        title_run.font.size = Pt(24)
        title_run.font.bold = True
    
    # 作者和日期
    if 'meta' in content_dict:
        meta_para = doc.add_paragraph()
        if 'author' in content_dict['meta']:
            meta_para.add_run(f"作者: {content_dict['meta']['author']}    ")
        if 'date' in content_dict['meta']:
            meta_para.add_run(f"日期: {content_dict['meta']['date']}")
        doc.add_paragraph()
    
    # 段落
    for para in content_dict.get('paragraphs', []):
        doc.add_paragraph(para)
    
    # 标题层级
    for heading in content_dict.get('headings', []):
        level = heading.get('level', 2)
        if level == 1:
            doc.add_heading(heading.get('text', ''), level=1)
        elif level == 2:
            doc.add_heading(heading.get('text', ''), level=2)
        else:
            doc.add_heading(heading.get('text', ''), level=3)
    
    # 列表
    for lst in content_dict.get('lists', []):
        items = lst.get('items', [])
        ordered = lst.get('ordered', False)
        for i, item in enumerate(items, 1):
            if ordered:
                doc.add_paragraph(f"{i}. {item}", style='List Number')
            else:
                doc.add_paragraph(item, style='List Bullet')
    
    # 表格
    for table_data in content_dict.get('tables', []):
        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])
        
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # 表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # 数据行
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
    
    # 图片
    for img in content_dict.get('images', []):
        try:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(img.get('path', ''), width=Inches(6))
        except Exception:
            pass
    
    # 保存文档
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    
    results['success'] = True
    results['message'] = f"""
✅ Word 文档创建成功！

📄 文件路径: {output_path}
📏 文件大小: {output_path.stat().st_size / 1024:.1f} KB

可直接在 Microsoft Word 或其他兼容软件中打开。
"""
    return results


def format_results(results: dict) -> str:
    """格式化结果"""
    return results.get('message', '❌ 创建失败')


def main():
    parser = argparse.ArgumentParser(
        description='Word Document Writer for Ruyi72',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python word_writer.py --title "报告" --content '{"paragraphs": ["内容1", "内容2"]}'
  python word_writer.py --output "report.docx" --data '{"title": "项目报告", "headings": [{"level": 1, "text": "概述"}]}'
        """
    )
    
    parser.add_argument('--output', '-o', default='document.docx', help='Output file path')
    parser.add_argument('--title', '-t', default='新文档', help='Document title')
    parser.add_argument('--content', '-c', default='{}', help='Content as JSON')
    parser.add_argument('--data', '-d', help='Full document data as JSON')
    
    args = parser.parse_args()
    
    if args.data:
        content_dict = json.loads(args.data)
    else:
        content_dict = json.loads(args.content)
    
    # 添加标题到内容
    if 'title' not in content_dict:
        content_dict['title'] = args.title
    
    results = create_word_document(content_dict, args.output)
    print(format_results(results))


if __name__ == '__main__':
    main()
